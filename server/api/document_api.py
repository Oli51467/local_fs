from fastapi import APIRouter, HTTPException
from pydantic import BaseModel
from typing import Optional, List, Callable, Awaitable, Tuple, Set
import logging
import hashlib
import pathlib
import re
import asyncio
import shutil
import uuid
from datetime import datetime
import tempfile
from typing import Dict, Any
from PIL import Image
import pypandoc
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as DOCX_RELATIONSHIP_TYPE
from service.clip_embedding_service import get_clip_embedding_service
from service.embedding_service import get_embedding_service
from service.faiss_service import FaissManager
from service.image_faiss_service import ImageFaissManager
from service.sqlite_service import SQLiteManager
from service.text_splitter_service import init_text_splitter_service, get_text_splitter_service
from config.config import ServerConfig, DatabaseConfig
from model.document_request_model import (
    FileUploadRequest,
    FileUploadResponse
)

class UpdateDocumentPathRequest(BaseModel):
    """更新文档路径请求模型"""
    old_path: str
    new_path: str
    is_folder: bool = False

class DeleteDocumentRequest(BaseModel):
    """删除文档请求模型"""
    file_path: str
    is_folder: bool = False

class UnmountDocumentRequest(BaseModel):
    """取消挂载文档请求模型"""
    file_path: str
    is_folder: bool = False

class UnmountDocumentResponse(BaseModel):
    """取消挂载文档响应模型"""
    status: str
    message: str
    unmounted_documents: int
    unmounted_vectors: int

class DeleteDocumentResponse(BaseModel):
    """删除文档响应模型"""
    status: str
    message: str
    deleted_documents: int
    deleted_vectors: int

class ReuploadDocumentRequest(BaseModel):
    """重新上传文档请求模型"""
    file_path: str
    force_reupload: bool = False  # 是否强制重新上传，忽略哈希检查

class ReuploadDocumentResponse(BaseModel):
    """重新上传文档响应模型"""
    status: str
    message: str
    document_id: Optional[int] = None
    file_info: Optional[Dict[str, Any]] = None


class FolderUploadStatusRequest(BaseModel):
    """文件夹上传状态请求模型"""
    folder_path: str


class FolderUploadStatusResponse(BaseModel):
    """文件夹上传状态响应模型"""
    folder: str
    files: Dict[str, bool]
    uploaded_files: List[str]


class FolderOperationRequest(BaseModel):
    """文件夹挂载/取消挂载请求"""
    folder_path: str


class FolderRemountRequest(BaseModel):
    """文件夹重新挂载请求"""
    folder_path: str
    force_reupload: bool = True

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api/document", tags=["document"])

# 全局变量
faiss_manager = None
sqlite_manager = None
text_splitter_service = None
image_faiss_manager = None

IMAGE_EXTENSIONS = {'.png', '.jpg', '.jpeg', '.gif', '.tiff', '.bmp'}
MARKDOWN_TYPES = {'md', 'markdown'}
DOCX_TYPES = {'docx'}
DOC_TYPES = {'doc'}
WORD_TYPES = DOC_TYPES.union(DOCX_TYPES)


def read_text_file_with_fallback(file_path: pathlib.Path) -> str:
    encodings = ['utf-8', 'utf-8-sig', 'gbk']
    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                return f.read()
        except UnicodeDecodeError:
            continue

    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
        return f.read()


def markdown_to_plain_text(markdown_text: str) -> str:
    text = markdown_text
    text = re.sub(r'^---[\s\S]*?---\s*', '', text, flags=re.MULTILINE)  # front matter
    text = re.sub(r'```[\s\S]*?```', '\n', text)
    text = re.sub(r'`([^`]+)`', r'\1', text)
    text = re.sub(r'!\[([^\]]*)\]\([^\)]+\)', r'\1', text)
    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
    text = re.sub(r'^#{1,6}\s*', '', text, flags=re.MULTILINE)
    text = re.sub(r'>\s?', '', text)
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
    text = re.sub(r'__([^_]+)__', r'\1', text)
    text = re.sub(r'\*([^*]+)\*', r'\1', text)
    text = re.sub(r'_([^_]+)_', r'\1', text)
    text = re.sub(r'~~([^~]+)~~', r'\1', text)
    text = re.sub(r'(?m)^\s*[-*+]\s+', '', text)
    text = re.sub(r'(?m)^\s*\d+\.\s+', '', text)
    text = re.sub(r'\s+\n', '\n', text)
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


def extract_text_content(file_path: pathlib.Path, file_type: str) -> str:
    lowered = file_type.lower()
    if lowered in MARKDOWN_TYPES:
        raw_text = read_text_file_with_fallback(file_path)
        return markdown_to_plain_text(raw_text)

    if lowered in DOCX_TYPES:
        text_content, _ = extract_docx_text_and_images(file_path, collect_images=False)
        return text_content

    if lowered in DOC_TYPES:
        text_content, _ = extract_doc_text_and_images(file_path, collect_images=False)
        return text_content

    raw_text = read_text_file_with_fallback(file_path)
    return raw_text


def _normalize_markdown_image_target(target: str) -> str:
    cleaned = (target or "").strip()
    if not cleaned:
        return ""

    if cleaned.startswith('<') and cleaned.endswith('>'):
        cleaned = cleaned[1:-1].strip()

    if cleaned.startswith('http://') or cleaned.startswith('https://'):
        return ""

    # 移除标题部分
    if '"' in cleaned:
        cleaned = cleaned.split('"', 1)[0].strip()
    elif "'" in cleaned:
        cleaned = cleaned.split("'", 1)[0].strip()

    cleaned = cleaned.strip()
    if not cleaned:
        return ""

    if re.search(r"\s", cleaned):
        cleaned = cleaned.split()[0]

    return cleaned


def _resolve_markdown_image_path(base_dir: pathlib.Path, target: str) -> Optional[pathlib.Path]:
    if not target:
        return None

    candidate = pathlib.Path(target)
    try:
        if candidate.is_absolute():
            resolved = candidate.resolve(strict=True)
        else:
            resolved = (base_dir / candidate).resolve(strict=True)
    except (FileNotFoundError, RuntimeError):
        return None

    if resolved.is_file():
        return resolved
    return None


def extract_markdown_images(file_path: pathlib.Path, markdown_text: str) -> List[Dict[str, Any]]:
    images: List[Dict[str, Any]] = []
    if not markdown_text:
        return images

    pattern = re.compile(r'!\[([^\]]*)\]\(([^)]+)\)')
    seen_paths: set[pathlib.Path] = set()

    for match in pattern.finditer(markdown_text):
        alt_text = match.group(1).strip()
        target_raw = match.group(2)
        normalized_target = _normalize_markdown_image_target(target_raw)
        if not normalized_target:
            continue

        resolved = _resolve_markdown_image_path(file_path.parent, normalized_target)
        if not resolved:
            logger.warning("Markdown图片路径无效或不存在: %s", normalized_target)
            continue

        if resolved.suffix.lower() not in IMAGE_EXTENSIONS:
            logger.debug("忽略不支持的图片格式: %s", resolved)
            continue

        resolved = resolved.resolve()
        if resolved in seen_paths:
            logger.debug("跳过重复的图片引用: %s", resolved)
            continue

        try:
            file_stat = resolved.stat()
        except FileNotFoundError:
            logger.warning("无法获取图片文件信息，文件不存在: %s", resolved)
            continue

        try:
            with Image.open(resolved) as image_obj:
                width, height = image_obj.size
        except Exception as exc:  # pylint: disable=broad-except
            logger.warning("读取图片尺寸失败 %s: %s", resolved, exc)
            continue

        line_number = markdown_text.count('\n', 0, match.start()) + 1

        try:
            relative_source = str(resolved.relative_to(ServerConfig.PROJECT_ROOT))
        except ValueError:
            relative_source = str(resolved)

        images.append({
            "source_path": resolved,
            "source_path_relative": relative_source,
            "line_number": line_number,
            "alt_text": alt_text,
            "image_format": resolved.suffix.lstrip('.').lower(),
            "image_size": file_stat.st_size,
            "width": width,
            "height": height,
            "markdown_target": normalized_target
        })

        seen_paths.add(resolved)

    return images


def _extract_docx_paragraph_text(document: Document) -> List[str]:
    lines: List[str] = []

    def append_if_content(value: Optional[str]) -> None:
        if value:
            stripped = value.strip()
            if stripped:
                lines.append(stripped)

    for paragraph in document.paragraphs:
        append_if_content(paragraph.text)

    for table in document.tables:
        for row in table.rows:
            cell_texts = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
            if cell_texts:
                lines.append('\t'.join(cell_texts))

    return lines


def _collect_docx_alt_text_map(document: Document) -> Dict[str, str]:
    alt_map: Dict[str, str] = {}
    for inline_shape in document.inline_shapes:
        inline = inline_shape._inline  # type: ignore[attr-defined]
        doc_pr = getattr(inline, 'docPr', None)
        alt_text = ''
        if doc_pr is not None and hasattr(doc_pr, 'attrib'):
            alt_text = (doc_pr.attrib.get('descr') or doc_pr.attrib.get('title') or '').strip()
        try:
            blip = inline.graphic.graphicData.pic.blipFill.blip  # type: ignore[attr-defined]
            embed_id = getattr(blip, 'embed', None)
            if embed_id:
                alt_map[embed_id] = alt_text
        except AttributeError:
            continue
    return alt_map


def extract_docx_text_and_images(file_path: pathlib.Path, collect_images: bool = True) -> Tuple[str, List[Dict[str, Any]]]:
    try:
        document = Document(str(file_path))
    except Exception as exc:  # pylint: disable=broad-except
        logger.error("读取DOCX文档失败 %s: %s", file_path, exc)
        raise HTTPException(status_code=400, detail=f"无法解析DOCX文档: {exc}") from exc

    paragraph_lines = _extract_docx_paragraph_text(document)
    text_content = '\n'.join(paragraph_lines).strip()

    if not collect_images:
        return text_content, []

    images: List[Dict[str, Any]] = []
    temp_dir: Optional[pathlib.Path] = None
    alt_map = _collect_docx_alt_text_map(document)

    try:
        for rel in document.part.rels.values():  # type: ignore[attr-defined]
            if rel.reltype != DOCX_RELATIONSHIP_TYPE.IMAGE:
                continue
            image_part = rel.target_part
            image_name = pathlib.Path(str(image_part.partname)).name
            suffix = pathlib.Path(image_name).suffix.lower()

            if suffix and suffix not in IMAGE_EXTENSIONS:
                # 跳过不支持的图片格式
                continue

            if temp_dir is None:
                temp_dir = pathlib.Path(tempfile.mkdtemp(prefix='docx_img_'))

            dest_path = temp_dir / image_name
            try:
                with open(dest_path, 'wb') as handle:
                    handle.write(image_part.blob)
            except Exception as exc:  # pylint: disable=broad-except
                logger.warning("提取DOCX图片失败 %s: %s", image_name, exc)
                continue

            width = height = None
            try:
                with Image.open(dest_path) as pil_image:
                    width, height = pil_image.size
            except Exception:  # pylint: disable=broad-except
                width = height = None

            stat_info = dest_path.stat()
            images.append({
                "source_path": dest_path,
                "source_path_relative": None,
                "line_number": None,
                "alt_text": alt_map.get(rel.rId, ''),
                "image_format": suffix.lstrip('.'),
                "image_size": stat_info.st_size,
                "width": width,
                "height": height,
                "docx_relationship_id": rel.rId,
                "temp_dir": temp_dir
            })
    except Exception as exc:  # pylint: disable=broad-except
        logger.warning("解析DOCX图片失败: %s", exc)

    return text_content, images


def convert_doc_to_docx(source_path: pathlib.Path) -> Tuple[pathlib.Path, pathlib.Path]:
    temp_dir = pathlib.Path(tempfile.mkdtemp(prefix='doc_convert_'))
    output_path = temp_dir / f"{source_path.stem}.docx"
    try:
        pypandoc.convert_file(str(source_path), 'docx', outputfile=str(output_path))
    except Exception as exc:  # pylint: disable=broad-except
        shutil.rmtree(temp_dir, ignore_errors=True)
        logger.error("转换DOC到DOCX失败 %s: %s", source_path, exc)
        raise HTTPException(status_code=400, detail=f"无法转换DOC文档: {exc}") from exc

    if not output_path.exists():
        shutil.rmtree(temp_dir, ignore_errors=True)
        raise HTTPException(status_code=400, detail="DOC文档转换失败，输出文件不存在")

    return output_path, temp_dir


def extract_doc_text_and_images(file_path: pathlib.Path, collect_images: bool = True) -> Tuple[str, List[Dict[str, Any]]]:
    converted_path, temp_dir = convert_doc_to_docx(file_path)
    try:
        text_content, images = extract_docx_text_and_images(converted_path, collect_images=collect_images)
    finally:
        try:
            converted_path.unlink(missing_ok=True)
        except Exception:  # pylint: disable=broad-except
            pass
        shutil.rmtree(temp_dir, ignore_errors=True)

    return text_content, images


def extract_text_and_images(file_path: pathlib.Path, file_type: str) -> Tuple[str, List[Dict[str, Any]]]:
    lowered = file_type.lower()

    if lowered in MARKDOWN_TYPES:
        raw_text = read_text_file_with_fallback(file_path)
        images = extract_markdown_images(file_path, raw_text)
        return markdown_to_plain_text(raw_text), images

    if lowered in DOCX_TYPES:
        return extract_docx_text_and_images(file_path, collect_images=True)

    if lowered in DOC_TYPES:
        return extract_doc_text_and_images(file_path)

    raw_text = read_text_file_with_fallback(file_path)
    return raw_text, []


def store_document_images(document_id: int, file_path: pathlib.Path, images: List[Dict[str, Any]]) -> Dict[str, Any]:
    result = {
        "stored": 0,
        "vector_ids": [],
        "folder": None,
    }

    if not images:
        return result

    if sqlite_manager is None or image_faiss_manager is None:
        logger.warning("图像处理依赖未初始化，跳过图片向量化")
        return result

    try:
        clip_service = get_clip_embedding_service()
    except Exception as exc:  # pylint: disable=broad-except
        logger.error("加载CLIP模型失败: %s", exc)
        raise HTTPException(status_code=500, detail=f"加载图片嵌入模型失败: {exc}") from exc

    project_root = ServerConfig.PROJECT_ROOT.resolve()
    images_root = DatabaseConfig.IMAGES_DIR

    # 创建唯一的图片存储文件夹
    dest_folder = None
    for _ in range(3):
        candidate = images_root / uuid.uuid4().hex
        try:
            candidate.mkdir(parents=True, exist_ok=False)
            dest_folder = candidate
            break
        except FileExistsError:
            continue

    if dest_folder is None:
        raise HTTPException(status_code=500, detail="创建图片存储目录失败")

    try:
        relative_document_path = str(file_path.resolve().relative_to(project_root))
    except ValueError:
        relative_document_path = str(file_path.resolve())

    relative_folder = str(dest_folder.relative_to(project_root))

    stored_records: List[Dict[str, Any]] = []
    vectors: List[List[float]] = []
    faiss_metadata: List[Dict[str, Any]] = []
    cleanup_dirs: Set[pathlib.Path] = set()

    try:
        for index, info in enumerate(images):
            temp_dir = info.get('temp_dir')
            if temp_dir:
                cleanup_dirs.add(pathlib.Path(temp_dir))

            source_path = info.get('source_path')
            if not source_path:
                logger.warning("图片源路径缺失，跳过当前图片")
                continue

            source_path_path = pathlib.Path(source_path)

            dest_name = f"{index:04d}_{source_path_path.name}"
            dest_path = dest_folder / dest_name

            try:
                shutil.copy2(source_path_path, dest_path)
            except Exception as exc:  # pylint: disable=broad-except
                logger.warning("复制图片失败 %s -> %s: %s", source_path_path, dest_path, exc)
                continue

            try:
                vector = clip_service.encode_image_path(dest_path)
            except Exception as exc:  # pylint: disable=broad-except
                logger.error("向量化图片失败 %s: %s", dest_path, exc)
                try:
                    dest_path.unlink()
                except OSError:
                    pass
                continue

            stored_record = {
                **info,
                "storage_path": dest_path,
                "storage_name": dest_name
            }
            stored_records.append(stored_record)
            vectors.append(vector)

            try:
                storage_path_rel = str(dest_path.relative_to(project_root))
            except ValueError:
                storage_path_rel = str(dest_path)

            faiss_metadata.append({
                "document_id": document_id,
                "file_path": relative_document_path,
                "storage_path": storage_path_rel,
                "storage_folder": relative_folder,
                "image_name": dest_name,
                "image_format": info['image_format'],
                "image_size": info['image_size'],
                "width": info['width'],
                "height": info['height'],
                "line_number": info.get('line_number'),
                "alt_text": info.get('alt_text'),
                "source_path": info.get('source_path_relative'),
            })

        if not vectors:
            shutil.rmtree(dest_folder, ignore_errors=True)
            return result

        import numpy as np

        try:
            vectors_array = np.array(vectors, dtype=np.float32)
            vector_ids = image_faiss_manager.add_vectors(vectors_array, faiss_metadata)
        except Exception as exc:  # pylint: disable=broad-except
            shutil.rmtree(dest_folder, ignore_errors=True)
            logger.error("保存图片向量失败: %s", exc)
            raise HTTPException(status_code=500, detail=f"保存图片向量失败: {exc}") from exc

        for record, vector_id in zip(stored_records, vector_ids):
            storage_path = record['storage_path']
            try:
                storage_path_rel = str(storage_path.relative_to(project_root))
            except ValueError:
                storage_path_rel = str(storage_path)

            sqlite_manager.insert_document_image(
                document_id=document_id,
                chunk_index=None,
                line_number=record.get('line_number'),
                image_name=record['storage_name'],
                image_format=record['image_format'],
                image_size=record['image_size'],
                width=record['width'],
                height=record['height'],
                storage_path=storage_path_rel,
                storage_folder=relative_folder,
                source_path=record.get('source_path_relative'),
                vector_id=vector_id
            )

        result.update({
            "stored": len(vector_ids),
            "vector_ids": vector_ids,
            "folder": relative_folder,
        })
        logger.info("成功处理 %d 张图片，存储目录: %s", result['stored'], relative_folder)
        return result
    finally:
        for temp_dir in cleanup_dirs:
            try:
                shutil.rmtree(temp_dir, ignore_errors=True)
            except Exception:  # pylint: disable=broad-except
                pass


def gather_image_cleanup_targets(relative_path: str, is_folder: bool) -> Tuple[List[int], List[pathlib.Path]]:
    if sqlite_manager is None:
        return [], []

    if is_folder:
        image_vector_ids = sqlite_manager.get_image_vector_ids_by_path_prefix(relative_path)
        folders = sqlite_manager.get_image_storage_folders_by_path_prefix(relative_path)
    else:
        image_vector_ids = sqlite_manager.get_image_vector_ids_by_path(relative_path)
        folders = sqlite_manager.get_image_storage_folders_by_path(relative_path)

    images_root = DatabaseConfig.IMAGES_DIR.resolve()
    folder_paths: List[pathlib.Path] = []

    for folder in folders:
        if not folder:
            continue
        candidate = pathlib.Path(folder)
        if not candidate.is_absolute():
            candidate = (ServerConfig.PROJECT_ROOT / candidate).resolve()
        else:
            candidate = candidate.resolve()

        try:
            candidate.relative_to(images_root)
        except ValueError:
            logger.warning("跳过非图片存储目录: %s", candidate)
            continue

        folder_paths.append(candidate)

    unique_folders = []
    seen = set()
    for folder in folder_paths:
        if folder not in seen:
            unique_folders.append(folder)
            seen.add(folder)

    return image_vector_ids, unique_folders


def remove_image_folders(folder_paths: List[pathlib.Path]) -> int:
    removed = 0
    for folder in folder_paths:
        try:
            if folder.exists() and folder.is_dir():
                shutil.rmtree(folder, ignore_errors=True)
                removed += 1
        except Exception as exc:  # pylint: disable=broad-except
            logger.warning("删除图片目录失败 %s: %s", folder, exc)
    return removed

def resolve_folder_path(folder_path: str) -> pathlib.Path:
    path_obj = pathlib.Path(folder_path)
    if not path_obj.is_absolute():
        path_obj = (ServerConfig.PROJECT_ROOT / path_obj).resolve()
    else:
        path_obj = path_obj.resolve()

    try:
        path_obj.relative_to(ServerConfig.PROJECT_ROOT)
    except ValueError:
        raise HTTPException(status_code=400, detail="文件夹必须位于项目根目录内")

    if not path_obj.exists() or not path_obj.is_dir():
        raise HTTPException(status_code=404, detail=f"文件夹不存在: {path_obj}")

    return path_obj


def collect_files_in_folder(folder_path: pathlib.Path) -> List[pathlib.Path]:
    files = []
    for child in folder_path.rglob('*'):
        if child.is_file() and not child.name.startswith('.'):
            files.append(child)
    return files


def normalize_project_relative_path(path_value: str) -> str:
    """将任意形式的路径转换为项目根目录下的相对路径"""
    trimmed = (path_value or "").strip()
    if not trimmed:
        raise HTTPException(status_code=400, detail="文件路径不能为空")

    project_root = ServerConfig.PROJECT_ROOT.resolve()
    candidate_path = pathlib.Path(trimmed)

    if candidate_path.is_absolute():
        resolved = candidate_path.resolve(strict=False)
    else:
        resolved = (project_root / candidate_path).resolve(strict=False)

    try:
        relative_path = resolved.relative_to(project_root)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail="文件必须位于项目根目录内") from exc

    return str(relative_path)


async def run_folder_tasks(
    files: List[pathlib.Path],
    operation: Callable[[pathlib.Path], Awaitable[Any]],
    max_concurrency: int = 8
) -> List[Dict[str, Any]]:
    semaphore = asyncio.Semaphore(max(1, min(max_concurrency, len(files))))
    project_root = ServerConfig.PROJECT_ROOT.resolve()
    results: List[Dict[str, Any]] = []

    def to_relative(path: pathlib.Path) -> Optional[str]:
        try:
            return str(path.resolve(strict=False).relative_to(project_root))
        except ValueError:
            return None

    async def worker(path: pathlib.Path) -> Dict[str, Any]:
        async with semaphore:
            relative_path = to_relative(path)
            try:
                outcome = await operation(path)
                status: Optional[str]
                success: bool
                if isinstance(outcome, dict):
                    status = outcome.get('status')
                else:
                    status = getattr(outcome, 'status', None)
                success = status not in {'error', 'failed'} if status else True
                return {
                    'path': str(path),
                    'relative_path': relative_path,
                    'status': status or 'success',
                    'detail': outcome.dict() if hasattr(outcome, 'dict') else outcome,
                    'success': success
                }
            except HTTPException as http_exc:
                logger.error("处理文件失败 (HTTP): %s - %s", path, http_exc.detail)
                return {
                    'path': str(path),
                    'relative_path': relative_path,
                    'status': 'error',
                    'detail': http_exc.detail,
                    'success': False
                }
            except Exception as exc:  # pylint: disable=broad-except
                logger.error("处理文件失败: %s - %s", path, exc)
                return {
                    'path': str(path),
                    'relative_path': relative_path,
                    'status': 'error',
                    'detail': str(exc),
                    'success': False
                }

    tasks = [worker(path) for path in files]
    results = await asyncio.gather(*tasks)
    return results

def init_document_api(faiss_mgr: FaissManager, sqlite_mgr: SQLiteManager, image_faiss_mgr: ImageFaissManager):
    """初始化文档API"""
    global faiss_manager, sqlite_manager, image_faiss_manager, text_splitter_service
    faiss_manager = faiss_mgr
    sqlite_manager = sqlite_mgr
    image_faiss_manager = image_faiss_mgr
    
    # 初始化文本分割服务
    if ServerConfig.TEXT_SPLITTER_TYPE == "recursive":
        init_text_splitter_service(
            "recursive",
            chunk_size=ServerConfig.RECURSIVE_CHUNK_SIZE,
            chunk_overlap=ServerConfig.RECURSIVE_CHUNK_OVERLAP,
            separators=ServerConfig.RECURSIVE_SEPARATORS
        )
    else:  # semantic
        init_text_splitter_service(
            "semantic",
            breakpoint_threshold_type=ServerConfig.SEMANTIC_BREAKPOINT_THRESHOLD_TYPE,
            breakpoint_threshold_amount=ServerConfig.SEMANTIC_BREAKPOINT_THRESHOLD_AMOUNT
        )
    
    text_splitter_service = get_text_splitter_service()
    logger.info(f"Document API initialized with {text_splitter_service.get_splitter_info()['type']} text splitter")


@router.post("/upload-status", response_model=FolderUploadStatusResponse)
async def get_folder_upload_status(request: FolderUploadStatusRequest):
    """查询指定文件夹下一层文件的上传状态"""
    if sqlite_manager is None:
        raise HTTPException(status_code=500, detail="数据库管理器未初始化")

    raw_path = request.folder_path.strip()
    if not raw_path or raw_path in {".", "./", ""}:
        normalized_path = "data"
    else:
        trimmed = raw_path.strip()
        trimmed = trimmed.lstrip("/")
        trimmed = trimmed.lstrip("./")
        if not trimmed:
            normalized_path = "data"
        elif trimmed.startswith("data"):
            normalized_path = trimmed.rstrip("/") or "data"
        else:
            normalized_path = f"data/{trimmed.rstrip('/')}"

    project_root = ServerConfig.PROJECT_ROOT.resolve()
    data_root = DatabaseConfig.DATABASE_DIR.resolve()

    folder_relative = pathlib.Path(normalized_path)
    folder_full_path = (project_root / folder_relative).resolve()

    try:
        folder_full_path.relative_to(data_root)
    except ValueError:
        raise HTTPException(status_code=400, detail="文件夹必须位于数据目录内")

    if not folder_full_path.exists() or not folder_full_path.is_dir():
        raise HTTPException(status_code=404, detail="指定的文件夹不存在")

    try:
        entries = [entry for entry in folder_full_path.iterdir() if entry.is_file()]
    except PermissionError as exc:
        logger.error("读取文件夹失败: %s", exc)
        raise HTTPException(status_code=500, detail="读取文件夹内容失败") from exc

    entries.sort(key=lambda path: path.name.lower())

    path_pairs = []
    for file_path in entries:
        try:
            rel_path = str(file_path.resolve().relative_to(project_root))
        except ValueError:
            logger.warning("文件 %s 不在项目根目录内，已跳过", file_path)
            rel_path = None
        path_pairs.append((file_path, rel_path))

    valid_rel_paths = [rel_path for _, rel_path in path_pairs if rel_path is not None]
    existing_paths = set(sqlite_manager.get_documents_by_paths(valid_rel_paths))

    files_status: Dict[str, bool] = {}
    for file_path, rel_path in path_pairs:
        files_status[file_path.name] = rel_path in existing_paths if rel_path else False

    uploaded_files = [name for name, uploaded in files_status.items() if uploaded]

    return FolderUploadStatusResponse(
        folder=str(folder_relative),
        files=files_status,
        uploaded_files=uploaded_files
    )

@router.post("/upload", response_model=FileUploadResponse)
async def upload_document(request: FileUploadRequest):
    """
    上传文档到系统
    
    Args:
        request: 文件上传请求，包含文件路径
        
    Returns:
        上传结果，包括状态、消息和文档信息
    """
    try:
        logger.info(f"收到文件上传请求: {request.file_path}")
        
        # 1. 验证文件路径
        file_path = pathlib.Path(request.file_path)
        if not file_path.exists():
            raise HTTPException(status_code=404, detail=f"文件不存在: {request.file_path}")
        
        if not file_path.is_file():
            raise HTTPException(status_code=400, detail=f"路径不是文件: {request.file_path}")
        
        # 2. 获取项目根目录
        project_root = ServerConfig.PROJECT_ROOT
        
        # 验证文件是否在项目根目录内
        try:
            file_path.relative_to(project_root)
        except ValueError:
            raise HTTPException(status_code=400, detail=f"文件必须在项目根目录内: {project_root}")
        
        # 获取文件名和类型
        filename = file_path.name
        file_type = file_path.suffix.lower().lstrip('.')
        logger.info(f"文件信息: 名称={filename}, 类型={file_type}")
        
        # 3. 计算文件哈希值检查是否已上传
        file_hash = calculate_file_hash(file_path)
        relative_file_path = str(file_path.relative_to(project_root))
        
        # 使用新的复合校验逻辑：同时检查文件路径和哈希
        if sqlite_manager:
            # 首先检查完全相同的文件路径和哈希（同一文件）
            existing_doc = sqlite_manager.get_document_by_path_and_hash(relative_file_path, file_hash)
            if existing_doc:
                # 检查是否已有块记录，如果没有则重新处理
                existing_chunks = sqlite_manager.get_document_chunks(existing_doc['id'])
                if not existing_chunks:
                    logger.info(f"文件已存在但无块记录，重新处理文档ID: {existing_doc['id']}")
                    # 继续执行后续的分割、嵌入等操作，但使用现有文档ID
                    document_id = existing_doc['id']
                    # 跳过文档插入步骤，继续执行后续步骤
                else:
                    return FileUploadResponse(
                        status="exists",
                        message=f"文件已存在，文档ID: {existing_doc['id']}",
                        document_id=existing_doc['id'],
                        file_info={
                            "filename": filename,
                            "file_type": file_type,
                            "file_hash": file_hash,
                            "file_size": file_path.stat().st_size
                        }
                    )
            else:
                # 检查是否有相同哈希但不同路径的文件（文件内容相同但位置不同）
                same_hash_docs = sqlite_manager.get_documents_by_hash(file_hash)
                if same_hash_docs:
                    # 文件内容相同但路径不同，可能是文件被移动了
                    logger.warning(f"发现相同哈希但不同路径的文件，可能是文件移动: {file_hash}")
                    
                    # 检查原文件是否还存在磁盘上
                    existing_doc = same_hash_docs[0]  # 取最新的一个
                    original_full_path = project_root / existing_doc['file_path']
                    
                    if original_full_path.exists():
                        # 原文件还存在，说明是不同位置的相同文件，拒绝上传
                        return FileUploadResponse(
                            status="exists",
                            message=f"相同内容的文件已存在于: {existing_doc['file_path']}",
                            document_id=existing_doc['id'],
                            file_info={
                                "filename": filename,
                                "file_type": file_type,
                                "file_hash": file_hash,
                                "file_size": file_path.stat().st_size,
                                "existing_path": existing_doc['file_path']
                            }
                        )
                    else:
                        # 原文件不存在，可能是文件被移动了，更新路径信息
                        logger.info(f"检测到文件移动，从 {existing_doc['file_path']} 到 {relative_file_path}")
                        
                        # 更新文档路径信息
                        sqlite_manager.update_document_path(existing_doc['file_path'], relative_file_path)
                        
                        # 更新Faiss向量元数据中的路径信息
                        if faiss_manager:
                            faiss_manager.update_metadata_by_path(existing_doc['file_path'], relative_file_path)
                        
                        return FileUploadResponse(
                            status="updated",
                            message=f"文件路径已更新: {existing_doc['file_path']} -> {relative_file_path}",
                            document_id=existing_doc['id'],
                            file_info={
                                "filename": filename,
                                "file_type": file_type,
                                "file_hash": file_hash,
                                "file_size": file_path.stat().st_size,
                                "old_path": existing_doc['file_path'],
                                "new_path": relative_file_path
                            }
                        )
        
        # 4. 根据文档类型提取文本
        supported_types = {"txt", "md", "markdown", "docx", "doc"}
        if file_type not in supported_types:
            raise HTTPException(status_code=400, detail=f"暂不支持的文件类型: {file_type}")

        text_content, extracted_images = extract_text_and_images(file_path, file_type)

        if extracted_images:
            logger.info("检测到 %d 张Markdown图片待处理", len(extracted_images))
        else:
            extracted_images = []

        if not text_content.strip():
            raise HTTPException(status_code=400, detail="文件内容为空")
        
        logger.info(f"文本提取完成，长度: {len(text_content)}")
        
        # 5. 分割文本
        if not text_splitter_service:
            raise HTTPException(status_code=500, detail="文本分割器未初始化")
        
        # 使用配置的文本分割器分割文本
        chunks = text_splitter_service.split_text(text_content)
        splitter_info = text_splitter_service.get_splitter_info()
        logger.info(f"文本分割完成，共 {len(chunks)} 个块，分割器类型: {splitter_info['type']}")
        
        if not chunks:
            raise HTTPException(status_code=400, detail="文本分割后无有效内容")
        
        # 6. 存储文档到SQLite数据库
        if not sqlite_manager:
            raise HTTPException(status_code=500, detail="数据库管理器未初始化")
        
        # 检查是否是重新处理的情况
        is_reprocessing = False
        if 'document_id' in locals():
            is_reprocessing = True
            logger.info(f"重新处理文档，文档ID: {document_id}")
            # 清理现有的块和向量数据
            relative_path_str = str(file_path.relative_to(project_root))

            existing_image_vector_ids, existing_folders = gather_image_cleanup_targets(relative_path_str, False)
            if existing_image_vector_ids and image_faiss_manager:
                image_faiss_manager.delete_vectors_by_ids(existing_image_vector_ids)
                logger.info("已删除旧的图片向量数量: %d", len(existing_image_vector_ids))
            if existing_folders:
                removed_count = remove_image_folders(existing_folders)
                logger.info("已清理旧的图片目录数量: %d", removed_count)

            sqlite_manager.delete_document_by_path(relative_path_str)
            # 重新插入文档记录
            document_id = sqlite_manager.insert_document(
                filename=filename,
                file_path=str(file_path.relative_to(project_root)),
                file_type=file_type,
                file_size=file_path.stat().st_size,
                file_hash=file_hash,
                content=text_content,
                metadata={
                    "upload_time": datetime.now().isoformat(),
                    "chunks_count": len(chunks),
                    "original_path": request.file_path
                }
            )
        else:
            # 新文档，正常插入
            document_id = sqlite_manager.insert_document(
                filename=filename,
                file_path=str(file_path.relative_to(project_root)),
                file_type=file_type,
                file_size=file_path.stat().st_size,
                file_hash=file_hash,
                content=text_content,
                metadata={
                    "upload_time": datetime.now().isoformat(),
                    "chunks_count": len(chunks),
                    "original_path": request.file_path
                }
            )
            logger.info(f"文档已存储到SQLite，文档ID: {document_id}")
        
        # 7. 生成文本嵌入向量
        embedding_service = get_embedding_service()
        if not embedding_service:
            raise HTTPException(status_code=500, detail="嵌入服务未初始化")
        
        # 为每个文本块生成嵌入向量
        embeddings = []
        for i, chunk in enumerate(chunks):
            try:
                embedding = embedding_service.encode_text(chunk)
                embeddings.append(embedding)
                logger.debug(f"已生成第 {i+1} 个文本块的嵌入向量")
            except Exception as e:
                logger.error(f"生成嵌入向量失败: {str(e)}")
                raise HTTPException(status_code=500, detail=f"生成嵌入向量失败: {str(e)}")
        
        logger.info(f"嵌入向量生成完成，共 {len(embeddings)} 个向量")
        
        # 8. 存储向量到Faiss索引
        if not faiss_manager:
            raise HTTPException(status_code=500, detail="Faiss管理器未初始化")
        
        # 为每个向量准备元数据
        vector_metadata = []
        for i, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
            metadata = {
                "document_id": document_id,
                "chunk_index": i,
                "chunk_text": chunk[:200] + "..." if len(chunk) > 200 else chunk,  # 存储前200字符作为预览
                "chunk_size": len(chunk),
                "file_path": str(file_path.relative_to(project_root)),
                "filename": filename,
                "file_type": file_type
            }
            vector_metadata.append(metadata)
        
        # 批量添加向量到索引
        import numpy as np
        embeddings_array = np.array(embeddings, dtype=np.float32)
        vector_ids = faiss_manager.add_vectors(embeddings_array, vector_metadata)
        logger.info(f"向量已存储到Faiss索引，向量ID列表: {vector_ids}")
        
        # 9. 存储文本块信息到数据库
        chunk_ids = []
        for i, (chunk, vector_id) in enumerate(zip(chunks, vector_ids)):
            chunk_id = sqlite_manager.insert_chunk(
                document_id=document_id,
                chunk_index=i,
                content=chunk,
                vector_id=vector_id
            )
            chunk_ids.append(chunk_id)
        
        logger.info(f"文本块信息已存储到数据库，块ID列表: {chunk_ids}")

        image_result = store_document_images(document_id, file_path, extracted_images)
        logger.info(
            "图片处理完成: 文档ID=%s, 图片数=%d",
            document_id,
            image_result.get('stored', 0)
        )
        
        # 10. 返回上传结果
        return FileUploadResponse(
            status="success",
            message=f"文档上传成功，文档ID: {document_id}",
            document_id=document_id,
            file_info={
                "filename": filename,
                "file_type": file_type,
                "file_hash": file_hash,
                "file_size": file_path.stat().st_size,
                "chunks_count": len(chunks),
                "vector_count": len(vector_ids),
                "image_count": image_result.get('stored', 0),
                "image_vector_count": len(image_result.get('vector_ids', [])),
                "image_folder": image_result.get('folder')
            }
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"文档上传失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"文档上传失败: {str(e)}")


@router.post("/mount-folder")
async def mount_folder(request: FolderOperationRequest) -> Dict[str, Any]:
    """批量挂载文件夹中的所有文件"""
    folder_path = resolve_folder_path(request.folder_path)
    files = collect_files_in_folder(folder_path)

    if not files:
        raise HTTPException(status_code=400, detail="文件夹内没有可挂载的文件")

    if len(files) > 50:
        raise HTTPException(status_code=400, detail="单次挂载的文件数量超过 50 个，请拆分后重试")

    logger.info("开始批量挂载文件夹: %s, 文件数量: %d", folder_path, len(files))

    async def mount_file(path: pathlib.Path):
        upload_request = FileUploadRequest(file_path=str(path))
        return await upload_document(upload_request)

    results = await run_folder_tasks(files, mount_file)

    success_count = sum(1 for item in results if item['success'])
    failure_count = len(results) - success_count

    status = 'success' if failure_count == 0 else ('partial' if success_count > 0 else 'failed')

    return {
        'status': status,
        'folder': str(folder_path),
        'total_files': len(results),
        'succeeded': success_count,
        'failed': failure_count,
        'details': results
    }


@router.post("/remount-folder")
async def remount_folder(request: FolderRemountRequest) -> Dict[str, Any]:
    """批量重新挂载文件夹中的所有文件"""
    folder_path = resolve_folder_path(request.folder_path)
    files = collect_files_in_folder(folder_path)

    if not files:
        raise HTTPException(status_code=400, detail="文件夹内没有可重新挂载的文件")

    if len(files) > 50:
        raise HTTPException(status_code=400, detail="单次重新挂载的文件数量超过 50 个，请拆分后重试")

    logger.info("开始批量重新挂载文件夹: %s, 文件数量: %d", folder_path, len(files))

    force = request.force_reupload
    project_root = ServerConfig.PROJECT_ROOT.resolve()

    async def remount_file(path: pathlib.Path):
        resolved_path = path.resolve()
        try:
            relative_path = str(resolved_path.relative_to(project_root))
        except ValueError as exc:
            raise HTTPException(status_code=400, detail="文件必须位于项目根目录内") from exc

        if not sqlite_manager:
            raise HTTPException(status_code=500, detail="数据库管理器未初始化")

        existing_doc = sqlite_manager.get_document_by_path(relative_path)

        if existing_doc:
            existing_hash = existing_doc.get('file_hash')
            current_hash = calculate_file_hash(resolved_path)

            if existing_hash == current_hash:
                logger.debug("文件未变化，跳过重新挂载: %s", relative_path)
                return {
                    'status': 'skipped',
                    'message': '文件内容未改变，跳过重新挂载',
                    'file_path': relative_path,
                    'file_hash': current_hash
                }

            reupload_request = ReuploadDocumentRequest(
                file_path=str(resolved_path),
                force_reupload=force
            )
            return await reupload_document(reupload_request)

        logger.debug("文件未挂载，执行首次挂载: %s", relative_path)
        upload_request = FileUploadRequest(file_path=str(resolved_path))
        return await upload_document(upload_request)

    results = await run_folder_tasks(files, remount_file)

    success_count = sum(1 for item in results if item['success'])
    failure_count = len(results) - success_count
    skipped_count = sum(1 for item in results if item.get('status') == 'skipped')
    status = 'success' if failure_count == 0 else ('partial' if success_count > 0 else 'failed')

    return {
        'status': status,
        'folder': str(folder_path),
        'total_files': len(results),
        'succeeded': success_count,
        'failed': failure_count,
        'skipped': skipped_count,
        'details': results
    }


@router.post("/unmount-folder")
async def unmount_folder(request: FolderOperationRequest) -> Dict[str, Any]:
    """批量取消挂载文件夹中的所有文件"""
    folder_path = resolve_folder_path(request.folder_path)
    files = collect_files_in_folder(folder_path)

    if not files:
        raise HTTPException(status_code=400, detail="文件夹内没有可取消挂载的文件")

    if len(files) > 50:
        raise HTTPException(status_code=400, detail="单次取消挂载的文件数量超过 50 个，请拆分后重试")

    logger.info("开始批量取消挂载文件夹: %s, 文件数量: %d", folder_path, len(files))

    async def unmount_file(path: pathlib.Path):
        unmount_request = UnmountDocumentRequest(file_path=str(path), is_folder=False)
        return await unmount_document(unmount_request)

    results = await run_folder_tasks(files, unmount_file)

    success_count = sum(1 for item in results if item['success'])
    failure_count = len(results) - success_count
    status = 'success' if failure_count == 0 else ('partial' if success_count > 0 else 'failed')

    return {
        'status': status,
        'folder': str(folder_path),
        'total_files': len(results),
        'succeeded': success_count,
        'failed': failure_count,
        'details': results
    }
@router.post("/update-path")
async def update_document_path(request: UpdateDocumentPathRequest):
    """
    更新文档路径
    
    Args:
        request: 更新路径请求，包含旧路径和新路径
        
    Returns:
        更新结果
    """
    try:
        logger.info(f"收到路径更新请求: {request.old_path} -> {request.new_path}")
        
        if not sqlite_manager:
            raise HTTPException(status_code=500, detail="数据库管理器未初始化")
        
        # 验证路径参数
        if not request.old_path or not request.new_path:
            raise HTTPException(status_code=400, detail="路径参数不能为空")
        
        # 标准化路径（移除前导斜杠）
        old_path = request.old_path.lstrip('/')
        new_path = request.new_path.lstrip('/')
        
        # 根据是否是文件夹选择不同的更新方法
        if request.is_folder:
            # 更新文件夹下所有文档的路径
            updated_count = sqlite_manager.update_documents_by_path_prefix(old_path, new_path)
            logger.info(f"文件夹路径更新完成，更新了 {updated_count} 个文档")
        else:
            # 更新单个文档路径
            updated_count = sqlite_manager.update_document_path(old_path, new_path)
            logger.info(f"文档路径更新完成，更新了 {updated_count} 个文档")

        updated_vectors = 0
        updated_image_vectors = 0
        if faiss_manager is None:
            logger.warning('Faiss 管理器未初始化，无法同步更新向量路径')
        else:
            try:
                if request.is_folder:
                    updated_vectors = faiss_manager.update_metadata_by_path_prefix(old_path, new_path)
                    logger.info(
                        "Faiss 元数据路径更新（文件夹）: %s -> %s，更新了 %d 条记录",
                        old_path,
                        new_path,
                        updated_vectors
                    )
                else:
                    updated_vectors = faiss_manager.update_metadata_by_path(old_path, new_path)
                    logger.info(
                        "Faiss 元数据路径更新（文件）: %s -> %s，更新了 %d 条记录",
                        old_path,
                        new_path,
                        updated_vectors
                    )
            except Exception as faiss_error:
                logger.error(
                    "更新 Faiss 元数据路径失败: %s -> %s，错误: %s",
                    old_path,
                    new_path,
                    faiss_error
                )

        if image_faiss_manager is None:
            logger.warning('图片 Faiss 管理器未初始化，无法同步更新图片向量路径')
        else:
            try:
                if request.is_folder:
                    updated_image_vectors = image_faiss_manager.update_metadata_by_path_prefix(old_path, new_path)
                else:
                    updated_image_vectors = image_faiss_manager.update_metadata_by_path(old_path, new_path)
            except Exception as image_faiss_error:  # pylint: disable=broad-except
                logger.error(
                    "更新图片 Faiss 元数据路径失败: %s -> %s，错误: %s",
                    old_path,
                    new_path,
                    image_faiss_error
                )

        if updated_count > 0 or updated_vectors > 0 or updated_image_vectors > 0:
            return {
                "status": "success",
                "message": (
                    f"成功更新文档 {updated_count} 个、文本向量 {updated_vectors} 个路径"
                    f"、图片向量 {updated_image_vectors} 个路径"
                ),
                "updated_documents": updated_count,
                "updated_vectors": updated_vectors
            }
        else:
            return {
                "status": "not_found",
                "message": "未找到需要更新路径的文档或向量",
                "updated_documents": 0,
                "updated_vectors": 0
            }
            
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"更新文档路径失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"更新文档路径失败: {str(e)}")

@router.delete("/delete", response_model=DeleteDocumentResponse)
async def delete_document(request: DeleteDocumentRequest):
    """
    删除文档及相关数据
    
    Args:
        request: 删除文档请求，包含文件路径和是否为文件夹
        
    Returns:
        删除结果，包括删除的文档数量和向量数量
    """
    try:
        logger.info(f"收到文档删除请求: {request.file_path}, 文件夹: {request.is_folder}")
        
        if not sqlite_manager or not faiss_manager:
            raise HTTPException(status_code=500, detail="数据库管理器未初始化")
        
        # 验证路径参数
        if not request.file_path:
            raise HTTPException(status_code=400, detail="文件路径不能为空")
        
        # 标准化路径为项目内相对路径
        file_path = normalize_project_relative_path(request.file_path)
        
        deleted_docs = 0
        deleted_vectors = 0
        deleted_image_vectors = 0
        removed_image_dirs = 0

        if request.is_folder:
            # 删除文件夹及其下所有文档
            logger.info(f"开始递归删除文件夹: {file_path}")
            
            # 1. 获取文件夹下所有文档的向量ID
            vector_ids = sqlite_manager.get_vector_ids_by_path_prefix(file_path)
            logger.info(f"找到 {len(vector_ids)} 个向量需要删除")

            image_vector_ids, image_folders = gather_image_cleanup_targets(file_path, True)
            logger.info("找到 %d 个图片向量需要删除", len(image_vector_ids))
            
            # 2. 从SQLite中删除文档及相关数据
            deleted_docs = sqlite_manager.delete_documents_by_path_prefix(file_path)
            logger.info(f"从SQLite中删除了 {deleted_docs} 个文档")
            
            # 3. 从Faiss中删除向量
            if vector_ids:
                deleted_vectors = faiss_manager.delete_vectors_by_ids(vector_ids)
                logger.info(f"从Faiss中删除了 {deleted_vectors} 个向量")

            if image_vector_ids and image_faiss_manager:
                deleted_image_vectors = image_faiss_manager.delete_vectors_by_ids(image_vector_ids)
                logger.info("从图片Faiss中删除了 %d 个向量", deleted_image_vectors)

            if image_folders:
                removed_image_dirs = remove_image_folders(image_folders)
                logger.info("删除图片目录数量: %d", removed_image_dirs)
            
        else:
            # 删除单个文档
            logger.info(f"开始删除单个文档: {file_path}")
            
            # 1. 获取文档的向量ID
            vector_ids = sqlite_manager.get_vector_ids_by_path(file_path)
            logger.info(f"找到 {len(vector_ids)} 个向量需要删除")

            image_vector_ids, image_folders = gather_image_cleanup_targets(file_path, False)
            logger.info("找到 %d 个图片向量需要删除", len(image_vector_ids))
            
            # 2. 从SQLite中删除文档及相关数据
            deleted_docs = sqlite_manager.delete_document_by_path(file_path)
            logger.info(f"从SQLite中删除了 {deleted_docs} 个文档")
            
            # 3. 从Faiss中删除向量
            if vector_ids:
                deleted_vectors = faiss_manager.delete_vectors_by_ids(vector_ids)
                logger.info(f"从Faiss中删除了 {deleted_vectors} 个向量")

            if image_vector_ids and image_faiss_manager:
                deleted_image_vectors = image_faiss_manager.delete_vectors_by_ids(image_vector_ids)
                logger.info("从图片Faiss中删除了 %d 个向量", deleted_image_vectors)

            if image_folders:
                removed_image_dirs = remove_image_folders(image_folders)
                logger.info("删除图片目录数量: %d", removed_image_dirs)

        return DeleteDocumentResponse(
            status="success",
            message=(
                f"成功删除了 {deleted_docs} 个文档、{deleted_vectors} 个文本向量"
                f" 和 {deleted_image_vectors} 个图片向量"
            ),
            deleted_documents=deleted_docs,
            deleted_vectors=deleted_vectors
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"删除文档失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"删除文档失败: {str(e)}")

@router.post("/unmount", response_model=UnmountDocumentResponse)
async def unmount_document(request: UnmountDocumentRequest):
    """
    取消挂载文档 - 删除数据库中的记录但不删除文件
    
    Args:
        request: 取消挂载请求，包含文件路径和是否为文件夹
        
    Returns:
        取消挂载结果，包括取消挂载的文档数量和向量数量
    """
    try:
        logger.info(f"收到文档取消挂载请求: {request.file_path}, 文件夹: {request.is_folder}")
        
        if not sqlite_manager or not faiss_manager:
            raise HTTPException(status_code=500, detail="数据库管理器未初始化")
        
        # 验证路径参数
        if not request.file_path:
            raise HTTPException(status_code=400, detail="文件路径不能为空")
        
        # 标准化路径为项目内相对路径
        file_path = normalize_project_relative_path(request.file_path)
        
        unmounted_docs = 0
        unmounted_vectors = 0
        unmounted_image_vectors = 0
        removed_image_dirs = 0

        if request.is_folder:
            # 取消挂载文件夹及其下所有文档
            logger.info(f"开始递归取消挂载文件夹: {file_path}")
            
            # 1. 获取文件夹下所有文档的向量ID
            vector_ids = sqlite_manager.get_vector_ids_by_path_prefix(file_path)
            logger.info(f"找到 {len(vector_ids)} 个向量需要取消挂载")

            image_vector_ids, image_folders = gather_image_cleanup_targets(file_path, True)
            logger.info("找到 %d 个图片向量需要取消挂载", len(image_vector_ids))
            
            # 2. 从SQLite中删除文档及相关数据（但不删除文件）
            unmounted_docs = sqlite_manager.delete_documents_by_path_prefix(file_path)
            logger.info(f"从SQLite中取消挂载了 {unmounted_docs} 个文档")
            
            # 3. 从Faiss中删除向量
            if vector_ids:
                unmounted_vectors = faiss_manager.delete_vectors_by_ids(vector_ids)
                logger.info(f"从Faiss中删除了 {unmounted_vectors} 个向量")

            if image_vector_ids and image_faiss_manager:
                unmounted_image_vectors = image_faiss_manager.delete_vectors_by_ids(image_vector_ids)
                logger.info("从图片Faiss中删除了 %d 个向量", unmounted_image_vectors)

            if image_folders:
                removed_image_dirs = remove_image_folders(image_folders)
                logger.info("删除图片目录数量: %d", removed_image_dirs)
            
        else:
            # 取消挂载单个文档
            logger.info(f"开始取消挂载单个文档: {file_path}")
            
            # 1. 获取文档的向量ID
            vector_ids = sqlite_manager.get_vector_ids_by_path(file_path)
            logger.info(f"找到 {len(vector_ids)} 个向量需要取消挂载")

            image_vector_ids, image_folders = gather_image_cleanup_targets(file_path, False)
            logger.info("找到 %d 个图片向量需要取消挂载", len(image_vector_ids))
            
            # 2. 从SQLite中删除文档及相关数据（但不删除文件）
            unmounted_docs = sqlite_manager.delete_document_by_path(file_path)
            logger.info(f"从SQLite中取消挂载了 {unmounted_docs} 个文档")
            
            # 3. 从Faiss中删除向量
            if vector_ids:
                unmounted_vectors = faiss_manager.delete_vectors_by_ids(vector_ids)
                logger.info(f"从Faiss中删除了 {unmounted_vectors} 个向量")

            if image_vector_ids and image_faiss_manager:
                unmounted_image_vectors = image_faiss_manager.delete_vectors_by_ids(image_vector_ids)
                logger.info("从图片Faiss中删除了 %d 个向量", unmounted_image_vectors)

            if image_folders:
                removed_image_dirs = remove_image_folders(image_folders)
                logger.info("删除图片目录数量: %d", removed_image_dirs)

        return UnmountDocumentResponse(
            status="success",
            message=(
                f"成功取消挂载了 {unmounted_docs} 个文档、{unmounted_vectors} 个文本向量"
                f" 和 {unmounted_image_vectors} 个图片向量"
            ),
            unmounted_documents=unmounted_docs,
            unmounted_vectors=unmounted_vectors
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"取消挂载文档失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"取消挂载文档失败: {str(e)}")

@router.post("/reupload", response_model=ReuploadDocumentResponse)
async def reupload_document(request: ReuploadDocumentRequest):
    """
    重新上传文档到系统
    
    Args:
        request: 重新上传请求，包含文件路径和是否强制重新上传
        
    Returns:
        重新上传结果，包括状态、消息和文档信息
    """
    try:
        logger.info(f"收到文件重新上传请求: {request.file_path}, 强制重新上传: {request.force_reupload}")
        
        # 1. 验证文件路径
        file_path = pathlib.Path(request.file_path)
        if not file_path.exists():
            raise HTTPException(status_code=404, detail=f"文件不存在: {request.file_path}")
        
        if not file_path.is_file():
            raise HTTPException(status_code=400, detail=f"路径不是文件: {request.file_path}")
        
        # 2. 获取项目根目录
        project_root = ServerConfig.PROJECT_ROOT
        
        # 验证文件是否在项目根目录内
        try:
            file_path.relative_to(project_root)
        except ValueError:
            raise HTTPException(status_code=400, detail=f"文件必须在项目根目录内: {project_root}")
        
        # 获取文件名和类型
        filename = file_path.name
        file_type = file_path.suffix.lower().lstrip('.')
        relative_file_path = str(file_path.relative_to(project_root))
        logger.info(f"文件信息: 名称={filename}, 类型={file_type}, 相对路径={relative_file_path}")
        
        # 3. 计算文件哈希值
        file_hash = calculate_file_hash(file_path)
        file_size = file_path.stat().st_size
        
        # 4. 检查文档是否已存在
        if not sqlite_manager:
            raise HTTPException(status_code=500, detail="数据库管理器未初始化")
        
        existing_doc = sqlite_manager.get_document_by_path(relative_file_path)
        
        if not existing_doc:
            # 文档不存在，直接上传
            logger.info(f"文档不存在，直接上传: {relative_file_path}")
            # 调用现有的上传逻辑
            upload_request = FileUploadRequest(file_path=request.file_path)
            upload_response = await upload_document(upload_request)
            
            return ReuploadDocumentResponse(
                status="uploaded",
                message="文档上传成功",
                document_id=upload_response.document_id,
                file_info=upload_response.file_info
            )
        
        # 文档已存在，检查是否需要重新上传
        existing_hash = existing_doc.get('file_hash')
        existing_document_id = existing_doc.get('id')
        
        # 如果哈希相同且没有强制重新上传，则不需要重新处理
        if existing_hash == file_hash and not request.force_reupload:
            logger.info(f"文件哈希相同，无需重新上传: {file_hash}")
            return ReuploadDocumentResponse(
                status="unchanged",
                message="文件内容未改变，无需重新上传",
                document_id=existing_document_id,
                file_info={
                    "filename": filename,
                    "file_type": file_type,
                    "file_hash": file_hash,
                    "file_size": file_size,
                    "existing_path": relative_file_path
                }
            )
        
        # 哈希不同或强制重新上传，删除旧数据并重新处理
        logger.info(f"文件哈希不同或强制重新上传，旧哈希: {existing_hash}, 新哈希: {file_hash}")
        
        # 5. 删除旧文档数据（包括文档、chunks和sqlite_sequence）
        logger.info(f"开始删除旧文档数据，文档ID: {existing_document_id}, 文件路径: {relative_file_path}")
        
        # 获取旧向量的ID（用于后续从Faiss删除）
        old_vector_ids = sqlite_manager.get_vector_ids_by_path(relative_file_path)
        logger.info(f"找到 {len(old_vector_ids)} 个旧向量需要删除")

        old_image_vector_ids, old_image_folders = gather_image_cleanup_targets(relative_file_path, False)
        logger.info("找到 %d 个旧图片向量需要删除", len(old_image_vector_ids))

        # 从SQLite中删除文档及相关数据（包括documents、document_chunks和sqlite_sequence）
        deleted_docs = sqlite_manager.delete_document_by_path(relative_file_path)
        if deleted_docs > 0:
            logger.info(f"成功从SQLite中删除文档及相关数据，删除了 {deleted_docs} 个文档记录")
            logger.info(f"同时删除了该文档的所有chunks数据，并重置了sqlite_sequence表")
        else:
            logger.warning(f"删除文档数据失败或文档不存在: {relative_file_path}")
        
        # 从Faiss中删除对应的向量
        deleted_vectors = 0
        if old_vector_ids and faiss_manager:
            deleted_vectors = faiss_manager.delete_vectors_by_ids(old_vector_ids)
            logger.info(f"从Faiss中删除了 {deleted_vectors} 个向量")

        if old_image_vector_ids and image_faiss_manager:
            deleted_image_vectors = image_faiss_manager.delete_vectors_by_ids(old_image_vector_ids)
            logger.info("从图片Faiss中删除了 %d 个向量", deleted_image_vectors)

        if old_image_folders:
            removed_image_dirs = remove_image_folders(old_image_folders)
            logger.info("删除旧图片目录数量: %d", removed_image_dirs)
        
        # 6. 重新上传文档（调用现有的上传逻辑）
        logger.info("开始重新上传文档")
        upload_request = FileUploadRequest(file_path=request.file_path)
        upload_response = await upload_document(upload_request)
        
        return ReuploadDocumentResponse(
            status="reuploaded",
            message="文档重新上传成功",
            document_id=upload_response.document_id,
            file_info=upload_response.file_info
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"文档重新上传失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"文档重新上传失败: {str(e)}")

def calculate_file_hash(file_path: pathlib.Path) -> str:
    """计算文件哈希值"""
    hash_sha256 = hashlib.sha256()
    with open(file_path, "rb") as f:
        for chunk in iter(lambda: f.read(4096), b""):
            hash_sha256.update(chunk)
    return hash_sha256.hexdigest()
